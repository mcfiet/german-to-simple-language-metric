"""
Build a labeled sentence dataset from Lebenshilfe simple/normal documents.

It extracts sentences from files under:
- data/ls-translations-lebenshilfe/ls/simple
- data/ls-translations-lebenshilfe/ls/normal

and writes a CSV with columns: sentence,label
"""

import argparse
import csv
import re
import zipfile
from pathlib import Path
from typing import Iterable, List, Tuple

from docx import Document
import xml.etree.ElementTree as ET


SUPPORTED_EXTS = {".docx", ".odt"}


def list_files(input_dir: Path) -> List[Path]:
    return sorted([p for p in input_dir.rglob("*") if p.is_file()])


def extract_blocks_docx(doc_path: Path) -> List[str]:
    doc = Document(doc_path)
    blocks: List[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    blocks.append(text)

    return blocks


def extract_blocks_odt(doc_path: Path) -> List[str]:
    blocks: List[str] = []
    with zipfile.ZipFile(doc_path) as zf:
        try:
            data = zf.read("content.xml")
        except KeyError:
            return blocks
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return blocks

    for elem in root.iter():
        tag = elem.tag
        if tag.endswith("}p") or tag.endswith("}h"):
            text = "".join(elem.itertext()).strip()
            if text:
                blocks.append(text)
    return blocks


def split_sentences(blocks: Iterable[str]) -> List[str]:
    text = " ".join(blocks)
    text = re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()
    if not text:
        return []
    raw_sentences = re.split(r"(?<=[.!?])\s+", text)
    sentences: List[str] = []
    for s in raw_sentences:
        s = s.strip()
        if s:
            sentences.append(s)
    return sentences


def extract_sentences(path: Path) -> List[str]:
    ext = path.suffix.lower()
    if ext == ".docx":
        blocks = extract_blocks_docx(path)
    elif ext == ".odt":
        blocks = extract_blocks_odt(path)
    else:
        return []
    return split_sentences(blocks)


def collect_labeled_sentences(input_dir: Path, label: int) -> Tuple[List[Tuple[str, int]], List[Path]]:
    rows: List[Tuple[str, int]] = []
    skipped: List[Path] = []
    for path in list_files(input_dir):
        if path.suffix.lower() not in SUPPORTED_EXTS:
            skipped.append(path)
            continue
        sentences = extract_sentences(path)
        for s in sentences:
            rows.append((s, label))
    return rows, skipped


def main():
    parser = argparse.ArgumentParser(
        description="Build labeled CSV from Lebenshilfe simple/normal documents."
    )
    parser.add_argument(
        "--simple-dir",
        type=Path,
        default=Path("data/ls-translations-lebenshilfe/ls/simple"),
        help="Directory with simple documents (docx/odt).",
    )
    parser.add_argument(
        "--normal-dir",
        type=Path,
        default=Path("data/ls-translations-lebenshilfe/ls/normal"),
        help="Directory with normal documents (docx/odt).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("data/ls_translations_labeled.csv"),
        help="Output CSV path (sentence,label).",
    )
    parser.add_argument(
        "--simple-label",
        type=int,
        default=1,
        help="Label value for simple sentences.",
    )
    parser.add_argument(
        "--normal-label",
        type=int,
        default=0,
        help="Label value for normal sentences.",
    )
    args = parser.parse_args()

    if not args.simple_dir.exists():
        raise FileNotFoundError(f"Missing simple dir: {args.simple_dir}")
    if not args.normal_dir.exists():
        raise FileNotFoundError(f"Missing normal dir: {args.normal_dir}")

    simple_rows, simple_skipped = collect_labeled_sentences(args.simple_dir, args.simple_label)
    normal_rows, normal_skipped = collect_labeled_sentences(args.normal_dir, args.normal_label)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    with args.output.open("w", encoding="utf-8", newline="") as f:
        writer = csv.writer(f)
        writer.writerow(["sentence", "label"])
        writer.writerows(simple_rows)
        writer.writerows(normal_rows)

    total = len(simple_rows) + len(normal_rows)
    print(f"[done] wrote {total} rows to {args.output}")
    print(f"[info] simple sentences: {len(simple_rows)}")
    print(f"[info] normal sentences: {len(normal_rows)}")
    if simple_skipped or normal_skipped:
        skipped = simple_skipped + normal_skipped
        print(f"[warn] skipped {len(skipped)} unsupported files (not .docx/.odt)")


if __name__ == "__main__":
    main()
