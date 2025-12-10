"""
Extract sentences from all DOCX files in the Lebenshilfe translations folder.

Each sentence is written on its own line in the output text file.
"""

import argparse
import re
from pathlib import Path
from typing import Iterable, List

from docx import Document


def list_docx_files(input_dir: Path) -> List[Path]:
    return sorted([p for p in input_dir.rglob("*.docx") if p.is_file()])


def extract_blocks(doc_path: Path) -> List[str]:
    doc = Document(doc_path)
    blocks: List[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    # Tables (collect all cell text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    blocks.append(text)

    return blocks


def split_sentences(blocks: Iterable[str]) -> List[str]:
    text = " ".join(blocks)
    text = re.sub(r"\s+", " ", text.replace("\xa0", " ")).strip()
    if not text:
        return []
    raw_sentences = re.split(r"(?<=[.!?])\s+", text)
    sentences: List[str] = []
    for s in raw_sentences:
        s = s.strip()
        if s:
            sentences.append(s)
    return sentences


def main():
    parser = argparse.ArgumentParser(
        description="Extract sentences from Lebenshilfe DOCX files."
    )
    parser.add_argument(
        "--input-dir",
        type=Path,
        default=Path("data/ls-translations-lebenshilfe/ls"),
        help="Directory containing DOCX files (searched recursively).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=Path("data/ls-translations-lebenshilfe/sentences.txt"),
        help="Where to write the extracted sentences.",
    )
    args = parser.parse_args()

    files = list_docx_files(args.input_dir)
    if not files:
        print(f"[warn] no DOCX files found under {args.input_dir}")
        return

    all_sentences: List[str] = []
    for path in files:
        blocks = extract_blocks(path)
        sentences = split_sentences(blocks)
        all_sentences.extend(sentences)
    args.output.parent.mkdir(parents=True, exist_ok=True)
    args.output.write_text("\n".join(all_sentences), encoding="utf-8")
    print(f"[done] wrote {len(all_sentences)} sentences from {len(files)} files to {args.output}")


if __name__ == "__main__":
    main()
